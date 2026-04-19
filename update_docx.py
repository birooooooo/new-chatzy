from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document('CHATZY_Research_Documentation.docx')

# ── helpers ───────────────────────────────────────────────────────────────────
def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_page_break(doc):
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    run = p.add_run()
    run._r.append(OxmlElement('w:br'))
    run._r[-1].set(qn('w:type'), 'page')

def blue_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x29, 0x97, 0xFF) if level == 1 else RGBColor(0x1A, 0x3A, 0x5C)
        run.font.size = Pt(18 if level == 1 else 14)
    return h

def sub_heading(doc, text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2C, 0x6E, 0xAB)
        run.font.size = Pt(12)
    return h

# ══════════════════════════════════════════════════════════════════════════════
# Find insertion point — just before Chapter 4 heading
# ══════════════════════════════════════════════════════════════════════════════
# We'll append new content at the end of the document, then re-save.
# Actually we insert a new page for the flowchart section after chapter 3.

# ── Find the paragraph index of "4. System Design" ──────────────────────────
target_idx = None
for i, p in enumerate(doc.paragraphs):
    if '4. System Design' in p.text or '4 System Design' in p.text:
        target_idx = i
        break

print(f'Chapter 4 found at paragraph index: {target_idx}')

# ══════════════════════════════════════════════════════════════════════════════
# Build new content to insert BEFORE chapter 4 (flowchart + screenshots)
# We insert paragraphs directly into the document XML at the right position
# ══════════════════════════════════════════════════════════════════════════════
from docx.oxml.ns import nsmap
from copy import deepcopy
from lxml import etree

body = doc.element.body

def make_page_break_elem():
    p_elem = OxmlElement('w:p')
    r_elem = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r_elem.append(br)
    p_elem.append(r_elem)
    return p_elem

def make_heading_elem(text, level=1, color_hex='2997FF', size_pt=18):
    p_elem = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), f'Heading{level}')
    pPr.append(pStyle)
    p_elem.append(pPr)
    r_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), color_hex)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size_pt * 2))
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(size_pt * 2))
    bold = OxmlElement('w:b')
    rPr.extend([bold, color, sz, szCs])
    r_elem.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    r_elem.append(t)
    p_elem.append(r_elem)
    return p_elem

def make_para_elem(text, bold=False, size_pt=11, color_hex=None,
                   align='left', space_after_pt=6, italic=False):
    p_elem = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), align)
    pPr.append(jc)
    spAfter = OxmlElement('w:spacing')
    spAfter.set(qn('w:after'), str(space_after_pt * 20))
    pPr.append(spAfter)
    p_elem.append(pPr)
    r_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        rPr.append(OxmlElement('w:b'))
    if italic:
        i_elem = OxmlElement('w:i')
        rPr.append(i_elem)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size_pt * 2))
    rPr.append(sz)
    if color_hex:
        color = OxmlElement('w:color')
        color.set(qn('w:val'), color_hex)
        rPr.append(color)
    r_elem.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_elem.append(t)
    p_elem.append(r_elem)
    return p_elem

def make_image_para_elem(doc_obj, img_path, width_inches=6.0):
    from docx.oxml.ns import qn
    p = doc_obj.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    return p._p

def make_screenshot_placeholder(doc_obj, label, width_inches=5.5, height_inches=3.5):
    """Make a grey placeholder box for a screenshot."""
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    fname = f'_ph_{label.replace(" ","_").replace("/","_")}.png'
    fig, ax = plt.subplots(figsize=(width_inches, height_inches))
    ax.set_facecolor('#1A1A2E')
    fig.patch.set_facecolor('#1A1A2E')
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis('off')
    # Phone frame
    rect = mpatches.FancyBboxPatch((3, 0.4), 4, 6.2,
                                   boxstyle='round,pad=0.3',
                                   facecolor='#111827', edgecolor='#2997FF', lw=2)
    ax.add_patch(rect)
    # Screen area
    screen = mpatches.FancyBboxPatch((3.3, 1.0), 3.4, 5.0,
                                     boxstyle='round,pad=0.1',
                                     facecolor='#0A0F1E', edgecolor='#444466', lw=1)
    ax.add_patch(screen)
    ax.text(5, 3.5, label, ha='center', va='center',
            color='#2997FF', fontsize=13, fontweight='bold', wrap=True)
    ax.text(5, 2.8, '[Insert Screenshot Here]', ha='center', va='center',
            color='#666688', fontsize=9)
    # Home button
    circle = plt.Circle((5, 0.75), 0.22, color='#333355', zorder=5)
    ax.add_patch(circle)
    ax.text(3.3, 6.35, '●', color='#333355', fontsize=6, va='center')
    plt.tight_layout(pad=0)
    plt.savefig(fname, dpi=120, bbox_inches='tight',
                facecolor='#1A1A2E')
    plt.close()
    p = doc_obj.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(fname, width=Inches(width_inches))
    return p._p, fname

# ══════════════════════════════════════════════════════════════════════════════
# Locate the chapter 4 paragraph element in the body
# ══════════════════════════════════════════════════════════════════════════════
chapter4_elem = None
if target_idx is not None:
    chapter4_elem = doc.paragraphs[target_idx]._p

# ── Collect elements to insert (in order) ────────────────────────────────────
to_insert = []

# PAGE BREAK before flowchart section
to_insert.append(make_page_break_elem())

# ── 3.5 Heading ───────────────────────────────────────────────────────────────
to_insert.append(make_heading_elem(
    '3.5  Application Flowchart', level=2,
    color_hex='1A3A5C', size_pt=14))

to_insert.append(make_para_elem(
    'The following flowchart illustrates the complete navigation and logic flow '
    'of the CHATZY application — from initial launch through authentication, '
    'the four main feature tabs, real-time messaging, AI chatbot, stories, and '
    'settings. Decision points are shown as diamonds; process screens as rounded '
    'rectangles; and data/storage operations in teal.',
    size_pt=11))

# Flowchart image
to_insert.append(make_image_para_elem(doc, 'chatzy_flowchart.png', width_inches=6.2))

to_insert.append(make_para_elem(
    'Figure 3.1 — CHATZY Application Flowchart',
    italic=True, size_pt=9, color_hex='888888', align='center', space_after_pt=4))

# PAGE BREAK before screenshots
to_insert.append(make_page_break_elem())

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4 SUPPLEMENT — Screenshots
# ══════════════════════════════════════════════════════════════════════════════
to_insert.append(make_heading_elem(
    '4.8  Application Screenshots', level=2,
    color_hex='1A3A5C', size_pt=14))

to_insert.append(make_para_elem(
    'The following screenshots demonstrate the key screens of the CHATZY '
    'application running on an Android device. The glass-morphism UI, dark '
    'theme, and animated elements are visible across all screens.',
    size_pt=11))

# Screenshot entries: (title, description, figure_number)
screenshots = [
    ('Splash Screen',
     'The animated splash screen displays the CHATZY logo with a fade-in '
     'animation. It checks for an existing authentication session before '
     'redirecting the user to either the login screen or the home screen.',
     '4.1'),
    ('Login Screen',
     'The login screen presents a glass-morphism form with email/username '
     'and password fields. Users can also tap "Forgot Password?" to trigger '
     'a Firebase password reset email.',
     '4.2'),
    ('Register Screen',
     'The registration screen collects the user\'s display name, unique '
     'username, email, password, phone number, and optional profile avatar. '
     'Username uniqueness is validated against Firestore before account creation.',
     '4.3'),
    ('Chats List Screen',
     'The chats list shows all conversations in real time with avatar, '
     'display name, last message preview, timestamp, and unread count badge. '
     'Pinned chats appear at the top.',
     '4.4'),
    ('Chat Screen — Private',
     'The private chat screen shows message bubbles with delivery/read '
     'status indicators, a typing indicator, and a media toolbar for '
     'sending images, audio, and files. Messages from the current user '
     'appear on the right in blue; received messages appear on the left.',
     '4.5'),
    ('AI Chatbot Screen',
     'The AI chatbot screen features the animated Huggy 3D character at '
     'the top whose expression changes based on sentiment analysis of the '
     'user\'s messages. The Gemini 2.0 Flash model generates contextual '
     'responses using the conversation history.',
     '4.6'),
    ('Stories Feed Screen',
     'The stories feed displays a horizontal carousel of active user stories '
     'at the top of the screen. Tapping a story ring opens the full-screen '
     'story viewer. The "+" button opens the story creation flow.',
     '4.7'),
    ('Settings Screen',
     'The settings screen provides navigation to profile editing, theme '
     'configuration (dark/light, Nebula/Deep Black background), notification '
     'preferences, privacy controls, AI preferences, and the about page.',
     '4.8'),
    ('Group Chat Screen',
     'The group chat screen supports multiple participants. Admins can manage '
     'members, update the group name and avatar, and remove participants. '
     'All messages are delivered in real time to all group members.',
     '4.9'),
    ('Theme Settings Screen',
     'Users can toggle between dark and light mode and select the background '
     'style. The Nebula style applies a gradient with star texture; Deep Black '
     'applies a pure black background for OLED power efficiency.',
     '4.10'),
]

ph_files = []
for idx, (title, desc, fig_num) in enumerate(screenshots):
    # Sub-heading for each screenshot
    to_insert.append(make_heading_elem(
        f'Figure {fig_num} — {title}', level=3,
        color_hex='2C6EAB', size_pt=12))

    # Placeholder image
    p_elem, fname = make_screenshot_placeholder(doc, title, width_inches=3.2, height_inches=5.5)
    ph_files.append(fname)
    to_insert.append(p_elem)

    # Caption
    to_insert.append(make_para_elem(
        f'Figure {fig_num}: {title}',
        italic=True, size_pt=9, color_hex='888888',
        align='center', space_after_pt=4))

    # Description
    to_insert.append(make_para_elem(desc, size_pt=11, space_after_pt=10))

    # Two screenshots side-by-side hint
    if idx < len(screenshots) - 1 and idx % 2 == 1:
        to_insert.append(make_page_break_elem())

# ══════════════════════════════════════════════════════════════════════════════
# Insert all elements before chapter 4
# ══════════════════════════════════════════════════════════════════════════════
if chapter4_elem is not None:
    for elem in reversed(to_insert):
        chapter4_elem.addprevious(elem)
    print(f'Inserted {len(to_insert)} elements before Chapter 4.')
else:
    # Append to end if chapter 4 not found
    for elem in to_insert:
        body.append(elem)
    print('Chapter 4 not found — appended to end.')

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save('CHATZY_Research_Documentation.docx')
print('Word document updated: CHATZY_Research_Documentation.docx')

# ── Cleanup placeholder PNGs ──────────────────────────────────────────────────
for f in ph_files:
    if os.path.exists(f):
        os.remove(f)
print('Placeholder images cleaned up.')
