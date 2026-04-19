"""
CHATZY – convert_to_docx.py
Generates a professional white-background academic thesis document.
Style: Calibri 11pt body, dark-navy blue headings, light-gray code blocks,
       dark-navy table headers with white text, alternating row shading.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# ── Colour palette ────────────────────────────────────────────────────────────
C_NAVY      = RGBColor(0x1A, 0x3A, 0x5C)   # dark navy  – H1 / table headers
C_BLUE      = RGBColor(0x2C, 0x6E, 0xAB)   # medium blue – H2
C_STEEL     = RGBColor(0x3A, 0x7C, 0xBD)   # steel blue  – H3
C_DARK_TEXT = RGBColor(0x22, 0x22, 0x22)   # near-black  – H4 / body
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_LIGHT_ROW = 'EEF4FF'   # alternating table row tint (hex string)
C_TBL_HDR   = '1A3A5C'   # table header fill (hex string)
C_CODE_BG   = 'F2F2F2'   # code block background (hex string)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.54)

# ── Default paragraph style ───────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = C_DARK_TEXT

# ── XML helpers ───────────────────────────────────────────────────────────────
def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """Set individual cell borders; kwargs: top/bottom/left/right = (val, sz, color)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, (val, sz, color) in kwargs.items():
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'),   val)
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_page_numbers(doc_obj):
    """Add centered page numbers to the footer of every section."""
    for section in doc_obj.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.clear()
        run = para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

def add_para(document, text='', bold=False, italic=False, size=11,
             color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6, line_spacing=None):
    p = document.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if line_spacing:
        p.paragraph_format.line_spacing = Pt(line_spacing)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return p

def add_hr(document, color='1A3A5C'):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── Read markdown ─────────────────────────────────────────────────────────────
with open('CHATZY_Research_Documentation.md', encoding='utf-8') as f:
    lines = f.readlines()

def strip_md(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*',     r'\1', text)
    text = re.sub(r'`(.+?)`',        r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'&nbsp;',          ' ',  text)
    return text.strip()

# ── Title page ────────────────────────────────────────────────────────────────
for _ in range(4):
    doc.add_paragraph()

# University name
p_uni = doc.add_paragraph()
p_uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_uni = p_uni.add_run('Cihan University — Erbil')
r_uni.font.name = 'Calibri'
r_uni.font.size = Pt(14)
r_uni.font.color.rgb = C_NAVY

p_dept = doc.add_paragraph()
p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_dept = p_dept.add_run('Department of Informatics and Software Engineering')
r_dept.font.name = 'Calibri'
r_dept.font.size = Pt(12)
r_dept.font.color.rgb = C_NAVY

for _ in range(2):
    doc.add_paragraph()

add_hr(doc)

for _ in range(2):
    doc.add_paragraph()

# App title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_title = p_title.add_run('CHATZY')
r_title.bold = True
r_title.font.name = 'Calibri'
r_title.font.size = Pt(32)
r_title.font.color.rgb = C_NAVY

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run('Real-Time Chat Application with AI Integration')
r_sub.bold = True
r_sub.font.name = 'Calibri'
r_sub.font.size = Pt(16)
r_sub.font.color.rgb = C_BLUE

doc.add_paragraph()

p_type = doc.add_paragraph()
p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_type = p_type.add_run('Final Year Project Report')
r_type.italic = True
r_type.font.name = 'Calibri'
r_type.font.size = Pt(13)
r_type.font.color.rgb = C_DARK_TEXT

for _ in range(2):
    doc.add_paragraph()

add_hr(doc)

for _ in range(2):
    doc.add_paragraph()

# Meta info table (clean, no grid border)
meta_tbl = doc.add_table(rows=6, cols=2)
meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
labels = ['Students', 'Supervisor', 'Department', 'Institution', 'Academic Year', 'Submission']
values = [
    'Birhat Tofiq   |   Hiwa Nihmat   |   Ahmad Qasim',
    'Ms. Duaa',
    'Informatics and Software Engineering',
    'Cihan University — Erbil',
    '2025 – 2026',
    'April 2026',
]
for row_idx, (lbl, val) in enumerate(zip(labels, values)):
    lbl_cell = meta_tbl.rows[row_idx].cells[0]
    val_cell = meta_tbl.rows[row_idx].cells[1]

    lbl_cell.width = Cm(4)
    lbl_p = lbl_cell.paragraphs[0]
    lbl_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lbl_r = lbl_p.add_run(lbl + ':')
    lbl_r.bold = True
    lbl_r.font.name = 'Calibri'
    lbl_r.font.size = Pt(11)
    lbl_r.font.color.rgb = C_NAVY

    val_p = val_cell.paragraphs[0]
    val_p.paragraph_format.left_indent = Cm(0.4)
    val_r = val_p.add_run(val)
    val_r.font.name = 'Calibri'
    val_r.font.size = Pt(11)
    val_r.font.color.rgb = C_DARK_TEXT

# Remove table borders
for row in meta_tbl.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'none')
            tcBorders.append(el)
        tcPr.append(tcBorders)

doc.add_page_break()

# ── Signature / Declaration page ─────────────────────────────────────────────
h_decl = doc.add_heading('Declaration', level=1)
for run in h_decl.runs:
    run.font.color.rgb = C_NAVY
    run.font.size = Pt(18)
    run.font.name = 'Calibri'

doc.add_paragraph(
    'We hereby declare that this project and the accompanying report are our own '
    'original work. All sources referenced and used in this report have been '
    'acknowledged. This work has not been submitted for any other academic qualification.'
)

doc.add_paragraph()

sig_tbl = doc.add_table(rows=4, cols=3)
sig_tbl.style = 'Table Grid'
for j, txt in enumerate(['Student', 'Signature', 'Date']):
    cell = sig_tbl.rows[0].cells[j]
    cell.text = txt
    shade_cell(cell, C_TBL_HDR)
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = C_WHITE
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

for i, name in enumerate(['Birhat Tofiq', 'Hiwa Nihmat', 'Ahmad Qasim'], 1):
    row = sig_tbl.rows[i].cells
    row[0].text = name
    row[1].text = '___________________________'
    row[2].text = '______________'
    shade_cell(row[0], C_LIGHT_ROW if i % 2 == 0 else 'FFFFFF')

doc.add_page_break()

# ── Add page numbers ──────────────────────────────────────────────────────────
add_page_numbers(doc)

# ── Markdown body renderer ────────────────────────────────────────────────────
in_code    = False
code_buf   = []
in_table   = False
table_rows = []

def flush_code(document, buf):
    if not buf:
        return
    # Code paragraph with gray background
    p = document.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run('\n'.join(buf))
    run.font.name  = 'Courier New'
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # Gray background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  C_CODE_BG)
    pPr.append(shd)

def render_table(document, rows):
    if len(rows) < 2:
        return
    header = [c.strip() for c in rows[0].strip('|').split('|')]
    data   = []
    for row in rows[2:]:
        if row.strip().startswith('|'):
            data.append([c.strip() for c in row.strip('|').split('|')])
    if not header:
        return
    ncols = len(header)
    nrows = 1 + len(data)
    tbl = document.add_table(rows=nrows, cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for j, txt in enumerate(header):
        cell = tbl.rows[0].cells[j]
        cell.text = ''
        shade_cell(cell, C_TBL_HDR)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(strip_md(txt))
        run.bold = True
        run.font.color.rgb = C_WHITE
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

    # Data rows
    for i, row in enumerate(data, 1):
        fill = C_LIGHT_ROW if i % 2 == 0 else 'FFFFFF'
        for j, txt in enumerate(row[:ncols]):
            cell = tbl.rows[i].cells[j]
            cell.text = ''
            shade_cell(cell, fill)
            p = cell.paragraphs[0]
            run = p.add_run(strip_md(txt))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)

    document.add_paragraph()

# ── Walk lines ────────────────────────────────────────────────────────────────
i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')

    # Code block
    if line.strip().startswith('```'):
        if not in_code:
            in_code  = True
            code_buf = []
        else:
            in_code = False
            flush_code(doc, code_buf)
            code_buf = []
        i += 1
        continue
    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # Table
    if line.strip().startswith('|'):
        table_rows.append(line)
        i += 1
        continue
    else:
        if table_rows:
            render_table(doc, table_rows)
            table_rows = []

    stripped = line.strip()

    # Blank / separator
    if stripped == '' or stripped == '---' or stripped == '&nbsp;':
        i += 1
        continue

    # Headings
    m = re.match(r'^(#{1,4})\s+(.*)', line)
    if m:
        level = len(m.group(1))
        text  = strip_md(m.group(2))
        # Skip the document title at the very top (already on title page)
        if level == 1 and 'CHATZY' in text and i < 5:
            i += 1
            continue
        h = doc.add_heading(text, level=min(level, 4))
        h.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
        h.paragraph_format.space_after  = Pt(4)
        for run in h.runs:
            run.font.name = 'Calibri'
            if level == 1:
                run.font.color.rgb = C_NAVY
                run.font.size = Pt(18)
                run.bold = True
            elif level == 2:
                run.font.color.rgb = C_BLUE
                run.font.size = Pt(14)
                run.bold = True
            elif level == 3:
                run.font.color.rgb = C_STEEL
                run.font.size = Pt(12)
                run.bold = True
            else:
                run.font.color.rgb = C_DARK_TEXT
                run.font.size = Pt(11)
                run.bold = True
        i += 1
        continue

    # Bullet list
    m = re.match(r'^(\s*)[*\-]\s+(.*)', line)
    if m:
        text = strip_md(m.group(2))
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.space_before = Pt(1)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        i += 1
        continue

    # Numbered list
    m = re.match(r'^\s*\d+\.\s+(.*)', line)
    if m:
        text = strip_md(m.group(1))
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.space_before = Pt(1)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        i += 1
        continue

    # Regular paragraph
    text = strip_md(line)
    if text:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(14)
        run = p.add_run(text)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(11)
        run.font.color.rgb = C_DARK_TEXT

    i += 1

# Flush any remaining buffers
if table_rows:
    render_table(doc, table_rows)
if code_buf:
    flush_code(doc, code_buf)

# ── Save ──────────────────────────────────────────────────────────────────────
out = 'CHATZY_Research_Documentation.docx'
doc.save(out)
print(f'Saved: {out}')
